"""Парсеры файлов разных форматов (без pandas)."""
import csv
import io
import json
import xml.etree.ElementTree as ET
from abc import ABC, abstractmethod


class ParsedData:
    """Результат парсинга файла."""
    def __init__(self, columns: list, sample_rows: list, dtypes: dict, separator: str = ",", all_rows: list = None):
        self.columns = columns
        self.sample_rows = sample_rows
        self.all_rows = all_rows or sample_rows
        self.dtypes = dtypes
        self.separator = separator

    def to_dict(self) -> dict:
        return {
            "columns": self.columns,
            "sample_rows": self.sample_rows,
            "dtypes": self.dtypes,
            "separator": self.separator,
        }


class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_bytes: bytes, filename: str = "") -> ParsedData:
        pass

    def _infer_dtypes(self, columns: list, rows: list) -> dict:
        result = {}
        for col in columns:
            val = None
            for row in rows:
                if col in row and row[col] not in (None, ""):
                    val = row[col]
                    break
            if val is None:
                result[col] = "string"
            elif isinstance(val, bool):
                result[col] = "boolean"
            elif isinstance(val, (int, float)):
                result[col] = "number"
            else:
                try:
                    float(val)
                    result[col] = "number"
                except (ValueError, TypeError):
                    result[col] = "string"
        return result


class CsvParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str = "") -> ParsedData:
        text = file_bytes.decode("utf-8-sig")
        sniffer = csv.Sniffer()
        try:
            dialect = sniffer.sniff(text[:2000])
            sep = dialect.delimiter
        except csv.Error:
            sep = ";"
        reader = csv.DictReader(io.StringIO(text), delimiter=sep)
        columns = reader.fieldnames or []
        rows = []
        for i, row in enumerate(reader):
            if i >= 3:
                break
            rows.append(dict(row))
        dtypes = self._infer_dtypes(columns, rows)
        return ParsedData(columns=columns, sample_rows=rows, dtypes=dtypes, separator=sep)


class XlsxParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str = "") -> ParsedData:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True)
        ws = wb.active
        all_rows = []
        columns = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                columns = [str(c) if c else f"col_{j}" for j, c in enumerate(row)]
            else:
                row_dict = {}
                for j, val in enumerate(row):
                    if j < len(columns):
                        row_dict[columns[j]] = val if val is not None else ""
                all_rows.append(row_dict)
        wb.close()
        dtypes = self._infer_dtypes(columns, all_rows[:3])
        return ParsedData(columns=columns, sample_rows=all_rows[:3], dtypes=dtypes, all_rows=all_rows)


class JsonParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str = "") -> ParsedData:
        data = json.loads(file_bytes.decode("utf-8-sig"))
        if isinstance(data, dict):
            data = [data]
        if not isinstance(data, list) or len(data) == 0:
            return ParsedData(columns=[], sample_rows=[], dtypes={})
        columns = list(data[0].keys())
        rows = data[:3]
        dtypes = self._infer_dtypes(columns, rows)
        return ParsedData(columns=columns, sample_rows=rows, dtypes=dtypes)


class XmlParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str = "") -> ParsedData:
        root = ET.fromstring(file_bytes.decode("utf-8-sig"))
        rows = []
        for child in root:
            row = {}
            for elem in child:
                row[elem.tag] = elem.text or ""
            if child.attrib:
                for k, v in child.attrib.items():
                    row[k] = v
            if row:
                rows.append(row)
        if not rows:
            return ParsedData(columns=[], sample_rows=[], dtypes={})
        columns = list(rows[0].keys())
        sample = rows[:3]
        dtypes = self._infer_dtypes(columns, sample)
        return ParsedData(columns=columns, sample_rows=sample, dtypes=dtypes)


class HtmlParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str = "") -> ParsedData:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(file_bytes.decode("utf-8-sig"), "html.parser")
        table = soup.find("table")
        if not table:
            return ParsedData(columns=[], sample_rows=[], dtypes={})
        header_row = table.find("tr")
        if not header_row:
            return ParsedData(columns=[], sample_rows=[], dtypes={})
        columns = [th.get_text(strip=True) for th in header_row.find_all(["th", "td"])]
        rows = []
        for tr in table.find_all("tr")[1:4]:
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            row = dict(zip(columns, cells))
            rows.append(row)
        dtypes = self._infer_dtypes(columns, rows)
        return ParsedData(columns=columns, sample_rows=rows, dtypes=dtypes)


class DocxParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str = "") -> ParsedData:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        # Ищем таблицы в документе
        if doc.tables:
            all_rows = []
            columns = []
            for table in doc.tables:
                header = [cell.text.strip() for cell in table.rows[0].cells]
                if not columns:
                    columns = header
                for row in table.rows[1:]:
                    row_dict = {}
                    for j, cell in enumerate(row.cells):
                        if j < len(columns):
                            row_dict[columns[j]] = cell.text.strip()
                    if any(v for v in row_dict.values()):
                        all_rows.append(row_dict)
            if all_rows:
                dtypes = self._infer_dtypes(columns, all_rows[:3])
                return ParsedData(columns=columns, sample_rows=all_rows[:3], dtypes=dtypes, all_rows=all_rows)
        # Если таблиц нет — ищем key:value паттерны
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not lines:
            return ParsedData(columns=[], sample_rows=[], dtypes={})
        # Пробуем key: value или key - value паттерны
        kv_rows = {}
        for line in lines:
            for sep_char in [":", " - ", "\t"]:
                if sep_char in line:
                    parts = line.split(sep_char, 1)
                    if len(parts) == 2 and len(parts[0].strip()) < 60:
                        kv_rows[parts[0].strip()] = parts[1].strip()
                        break
        if len(kv_rows) >= 2:
            columns = list(kv_rows.keys())
            row = kv_rows
            dtypes = self._infer_dtypes(columns, [row])
            return ParsedData(columns=columns, sample_rows=[row], dtypes=dtypes, all_rows=[row])
        # Пробуем разделить по табуляции или точке с запятой
        sep = "\t" if "\t" in lines[0] else ";"
        parts = lines[0].split(sep)
        if len(parts) > 1:
            columns = parts
            all_rows = []
            for line in lines[1:]:
                vals = line.split(sep)
                row_dict = dict(zip(columns, vals))
                all_rows.append(row_dict)
            dtypes = self._infer_dtypes(columns, all_rows[:3])
            return ParsedData(columns=columns, sample_rows=all_rows[:3], dtypes=dtypes, separator=sep, all_rows=all_rows)
        # Просто текст — одна колонка
        columns = ["text"]
        all_rows = [{"text": line} for line in lines]
        dtypes = {"text": "string"}
        return ParsedData(columns=columns, sample_rows=all_rows[:3], dtypes=dtypes, all_rows=all_rows)


class PdfParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str = "") -> ParsedData:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        all_text = ""
        for page in reader.pages:
            all_text += (page.extract_text() or "") + "\n"
        lines = [l.strip() for l in all_text.split("\n") if l.strip()]
        if not lines:
            return ParsedData(columns=[], sample_rows=[], dtypes={})
        # Пробуем найти табличные данные по разделителю
        for sep in ["\t", ";", ","]:
            parts = lines[0].split(sep)
            if len(parts) > 1:
                columns = [p.strip() for p in parts]
                all_rows = []
                for line in lines[1:]:
                    vals = [v.strip() for v in line.split(sep)]
                    row_dict = dict(zip(columns, vals))
                    all_rows.append(row_dict)
                dtypes = self._infer_dtypes(columns, all_rows[:3])
                return ParsedData(columns=columns, sample_rows=all_rows[:3], dtypes=dtypes, separator=sep, all_rows=all_rows)
        # Просто текст
        columns = ["text"]
        all_rows = [{"text": line} for line in lines]
        dtypes = {"text": "string"}
        return ParsedData(columns=columns, sample_rows=all_rows[:3], dtypes=dtypes, all_rows=all_rows)


PARSERS = {
    ".csv": CsvParser,
    ".xls": XlsxParser,
    ".xlsx": XlsxParser,
    ".json": JsonParser,
    ".xml": XmlParser,
    ".html": HtmlParser,
    ".htm": HtmlParser,
    ".docx": DocxParser,
    ".pdf": PdfParser,
}


def get_parser(file_extension: str) -> BaseParser:
    ext = file_extension.lower()
    parser_class = PARSERS.get(ext)
    if not parser_class:
        raise ValueError(f"Формат {ext} не поддерживается. Доступные: {list(PARSERS.keys())}")
    return parser_class()


SUPPORTED_FORMATS = list(PARSERS.keys())
